"""Turn an uploaded résumé, bio or deal sheet into plain text.

Deliberately narrow. Everything a sender wants to prove about themselves arrives as a PDF, a
Word file, or plain text, and each of those has one obvious way to read it. Anything else is
rejected by name rather than half-parsed into garbage that then becomes a "proof point" in an
email to a real investor.

Scanned PDFs are the one case worth calling out: they contain images of text, not text, so
they extract to almost nothing. That is reported as a specific, actionable message instead of
an empty document that silently contributes no proof points.
"""

from __future__ import annotations

import io
import re

MAX_UPLOAD_BYTES = 12 * 1024 * 1024
# Past this the document is stored truncated. A 25-page deal book adds nothing beyond its first
# few thousand words for this purpose, and the whole profile has to fit in a draft prompt.
MAX_TEXT_CHARS = 60_000
# Below this, a PDF is almost certainly scanned images rather than text.
MIN_MEANINGFUL_CHARS = 120

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


class DocumentTextError(Exception):
    pass


def _tidy(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\x00", "")
    # Collapse the run of blank lines PDF extraction leaves between every block, but keep
    # paragraph breaks: bullet structure is most of what makes a résumé readable.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned in requirements
        raise DocumentTextError("PDF support is not installed on the server (pypdf)") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise DocumentTextError(f"This PDF could not be opened: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:
            raise DocumentTextError("This PDF is password protected") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # One unreadable page should not lose the other nineteen.
            continue
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is pinned in requirements
        raise DocumentTextError("Word support is not installed on the server (python-docx)") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentTextError(
            f"This Word file could not be opened: {exc}. If it is a .doc, save it as .docx first."
        ) from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    # Career history and deal lists are very often tables, and dropping them would lose exactly
    # the quantified rows that make the best proof points.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _from_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extension_of(filename: str) -> str:
    name = (filename or "").strip().lower()
    dot = name.rfind(".")
    return name[dot:] if dot > 0 else ""


def extract_text(filename: str, data: bytes) -> tuple[str, bool]:
    """(text, truncated) for one uploaded file. Raises DocumentTextError with a usable message."""
    if not data:
        raise DocumentTextError("The file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentTextError(
            f"The file is {len(data) / 1024 / 1024:.1f} MB; the limit is "
            f"{MAX_UPLOAD_BYTES // 1024 // 1024} MB"
        )

    extension = extension_of(filename)
    if extension == ".pdf":
        raw = _from_pdf(data)
    elif extension == ".docx":
        raw = _from_docx(data)
    elif extension in (".txt", ".md"):
        raw = _from_plain(data)
    elif extension == ".doc":
        raise DocumentTextError("Old .doc files are not supported — save it as .docx and try again")
    else:
        raise DocumentTextError(
            f"{extension or 'That file type'} is not supported. Upload a PDF, DOCX, TXT or MD file."
        )

    text = _tidy(raw)
    if len(text) < MIN_MEANINGFUL_CHARS:
        if extension == ".pdf":
            raise DocumentTextError(
                "Almost no text came out of this PDF, which usually means it is a scan rather "
                "than a text document. Upload the Word original if you have one."
            )
        raise DocumentTextError("There was not enough readable text in this file to use")

    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS], True
    return text, False


def guess_kind(filename: str, text: str) -> str:
    """A first guess at what was uploaded, so the list is readable without hand-labelling."""
    name = (filename or "").lower()
    head = (text or "")[:2500].lower()

    if "resume" in name or "résumé" in name or "cv" in name.replace(".", " ").split():
        return "resume"
    if "bio" in name:
        return "bio"
    if "case" in name and "stud" in name:
        return "case_study"
    if "deal" in name or "tombstone" in name or "track record" in name:
        return "deal_sheet"

    if any(marker in head for marker in ("professional experience", "work experience", "employment history")):
        return "resume"
    if any(marker in head for marker in ("transaction", "acquisitions", "deal sheet")):
        return "deal_sheet"
    return "other"
